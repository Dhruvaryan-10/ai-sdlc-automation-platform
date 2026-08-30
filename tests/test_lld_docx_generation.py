"""LLD.docx export (Phase 4 test 13)."""

from docx import Document

from app.document_generator.brd_generator import generate_lld_docx

SAMPLE_MD = """# Sample — Low-Level Design

**Version:** 1
**Client:** Acme Corp

## 1. Introduction and Source Traceability
Detailed design for the sample system.

## 3. Classes, Responsibilities and Interfaces
- RegistrationService: creates and validates customer accounts.
- AccountRepository: persists account records.

## 4. Services and API Specifications
| Endpoint | Method | Purpose |
| /api/register | POST | Create a customer account |
"""


def test_generate_lld_docx_writes_readable_file(tmp_path):
    out = tmp_path / "proj" / "lld" / "LLD_v1.docx"
    result = generate_lld_docx(SAMPLE_MD, out)

    assert result == out
    assert out.exists()

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "1. Introduction and Source Traceability" in text
    assert any("RegistrationService" in p.text for p in doc.paragraphs)
    bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert bullet_paras, "class list should render as bullets"
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "Endpoint"


def test_lld_docx_matches_selected_content(tmp_path):
    """What gets written is exactly the content handed in (the 'selected' version)."""
    out = tmp_path / "LLD_v2.docx"
    marker = "UNIQUE-MARKER-lld-456"
    generate_lld_docx(SAMPLE_MD + f"\n\n## 14. Notes\n{marker}\n", out)

    doc = Document(str(out))
    assert any(marker in p.text for p in doc.paragraphs)
