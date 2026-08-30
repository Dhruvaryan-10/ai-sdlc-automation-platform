"""UserStories.docx export (Phase 3 test 9)."""

from docx import Document

from app.document_generator.brd_generator import generate_user_stories_docx

SAMPLE_MD = """# Sample — Draft User Stories

**Version:** 1
**Client:** Acme Corp

## US-001 — Customer Registration

**User Story:**
As a customer,
I want to create an account,
so that I can access the application.

**Acceptance Criteria:**
- Required registration information can be entered.
- Invalid registration information is rejected.

**Priority:** High
"""


def test_generate_user_stories_docx_writes_readable_file(tmp_path):
    out = tmp_path / "proj" / "user_stories" / "UserStories_v1.docx"
    result = generate_user_stories_docx(SAMPLE_MD, out)

    assert result == out
    assert out.exists()

    doc = Document(str(out))
    paragraphs = [p.text for p in doc.paragraphs]
    text = "\n".join(paragraphs)

    assert "US-001 — Customer Registration" in text
    assert any("Required registration information" in p.text for p in doc.paragraphs)
    bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert bullet_paras, "acceptance criteria should render as bullets"


def test_user_stories_docx_matches_selected_content(tmp_path):
    """What gets written is exactly the content handed in (the 'selected' version)."""
    out = tmp_path / "UserStories_v2.docx"
    marker = "UNIQUE-MARKER-xyz789"
    generate_user_stories_docx(SAMPLE_MD + f"\n\n## US-002 — Notes\n{marker}\n", out)

    doc = Document(str(out))
    assert any(marker in p.text for p in doc.paragraphs)
