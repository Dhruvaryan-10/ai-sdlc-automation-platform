"""Test Case DOCX export (Phase 6 item 32)."""

from docx import Document

from app.document_generator.brd_generator import generate_test_cases_docx

SAMPLE_MD = """# Sample — Test Cases

**Version:** 2
**Source:** Artifact-refined
**Built From:** BRD v2, HLD v1, LLD unavailable, User Stories v3
**Client:** Acme Corp
**Project Type:** Web Application

## TC-001 — Register a new customer account

**Requirement / User Story Reference:** FR-1
**BRD Reference:** FR-1
**User Story Reference:** US-001
**Priority:** High
**Test Type:** Functional

**Preconditions:**
- The registration page is reachable.

**Test Steps:**
- 1. Open the registration page.
- 2. Enter valid details and submit.

**Expected Result:**
A new account is created and the customer is signed in.
"""


def test_generate_test_cases_docx_writes_readable_file(tmp_path):
    out = tmp_path / "proj" / "test_cases" / "TestCases_v2.docx"
    result = generate_test_cases_docx(SAMPLE_MD, out)

    assert result == out
    assert out.exists()

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "TC-001 — Register a new customer account" in text
    # provenance header is preserved in the export
    assert any("Version: 2" in p.text for p in doc.paragraphs)
    assert any("Source: Artifact-refined" in p.text for p in doc.paragraphs)
    assert any("Built From: BRD v2, HLD v1" in p.text for p in doc.paragraphs)
    bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert bullet_paras, "preconditions / steps should render as bullets"


def test_test_cases_docx_matches_supplied_content(tmp_path):
    out = tmp_path / "TestCases_v3.docx"
    marker = "UNIQUE-MARKER-tc-789"
    generate_test_cases_docx(SAMPLE_MD + f"\n\n## TC-002 — Notes\n**Expected Result:** {marker}\n", out)

    doc = Document(str(out))
    assert any(marker in p.text for p in doc.paragraphs)
