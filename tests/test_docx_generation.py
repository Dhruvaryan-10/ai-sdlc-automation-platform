"""DOCX export (TEST 9). Both BRD and HLD exports must produce a real .docx."""

from docx import Document

from app.document_generator.brd_generator import generate_brd_docx, generate_hld_docx

SAMPLE_MD = """# Sample — High-Level Design

**Version:** 1
**Client:** Acme Corp

## 1. Architecture Overview
A layered web application.

## 2. Components
- API gateway
- Service layer
- Data store

| Component | Responsibility |
| API | Routing |
"""


def test_generate_brd_docx_writes_readable_file(tmp_path):
    out = tmp_path / "proj" / "BRD_v1.docx"
    result = generate_brd_docx(SAMPLE_MD, out)

    assert result == out
    assert out.exists()
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Architecture Overview" in text
    assert any("web application" in p.text for p in doc.paragraphs)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Sample — High-Level Design" in headings or "Architecture Overview" in headings


def test_generate_hld_docx_writes_readable_file(tmp_path):
    out = tmp_path / "proj" / "hld" / "HLD_v1.docx"
    result = generate_hld_docx(SAMPLE_MD, out)

    assert result == out
    assert out.exists()
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Architecture Overview" in text
    # the pipe table is rendered as a table, not paragraphs
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "Component"


def test_hld_export_matches_selected_content(tmp_path):
    """What gets written is exactly the content handed in (the 'selected' version)."""
    out = tmp_path / "HLD_v2.docx"
    marker = "UNIQUE-MARKER-abc123"
    generate_hld_docx(SAMPLE_MD + f"\n\n## 3. Notes\n{marker}\n", out)

    doc = Document(str(out))
    assert any(marker in p.text for p in doc.paragraphs)
